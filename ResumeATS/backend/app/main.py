from __future__ import annotations

import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse
from uuid import uuid4

import requests
from bs4 import BeautifulSoup
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from striprtf.striprtf import rtf_to_text

BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

ALLOWED_FORMATS = {"all", "docx", "pdf", "rtf"}

app = FastAPI(title="Resume ATS Refactor API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class GenerateResponse(BaseModel):
    job_title: str
    company: str
    preview: str
    thumbnail: str
    files: dict[str, str]


def _slugify(value: str) -> str:
    value = re.sub(r"[^a-zA-Z0-9\s-]", "", value).strip().lower()
    return re.sub(r"\s+", "-", value)[:40] or "resume"


def _validate_job_url(url: str) -> None:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise HTTPException(status_code=400, detail="job_url must be a valid http/https URL.")


def _extract_job(url: str) -> tuple[str, str, str]:
    _validate_job_url(url)
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Unable to fetch job URL: {exc}") from exc

    soup = BeautifulSoup(response.text, "html.parser")
    title = (soup.title.string if soup.title and soup.title.string else "Job Opportunity").strip()

    headings = [h.get_text(" ", strip=True) for h in soup.select("h1, h2")]
    body_text = " ".join(p.get_text(" ", strip=True) for p in soup.select("p"))
    description = " ".join(headings + [body_text])[:4500]

    company_meta = soup.find("meta", attrs={"property": "og:site_name"})
    company = company_meta.get("content", "Unknown Company") if company_meta else "Unknown Company"
    return title, company, description


def _read_resume(file: UploadFile, raw_bytes: bytes) -> tuple[str, Optional[Document]]:
    suffix = Path(file.filename or "").suffix.lower()
    if suffix == ".docx":
        doc = Document(file.file)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return text, doc
    if suffix == ".rtf":
        return rtf_to_text(raw_bytes.decode("utf-8", errors="ignore")), None
    text = raw_bytes.decode("utf-8", errors="ignore")
    return text, None


def _extract_top_keywords(job_desc: str, n: int = 12) -> list[str]:
    words = re.findall(r"\b[A-Za-z][A-Za-z+/#.-]{3,}\b", job_desc)
    stop = {
        "with", "from", "your", "that", "have", "will", "this", "about", "their", "which", "would", "role",
        "team", "years", "experience", "skills", "work", "ability", "using", "build", "across", "including",
    }
    counts = Counter(w.lower() for w in words if w.lower() not in stop)
    return [w for w, _ in counts.most_common(n)]


def _build_bullets(resume_lines: list[str], keywords: list[str]) -> list[str]:
    candidates = [ln.strip("•- ") for ln in resume_lines if len(ln.strip()) > 25]
    selected = candidates[:3] if candidates else ["Delivered measurable business value across core initiatives."]
    enriched = []
    for idx, bullet in enumerate(selected, start=1):
        kw = keywords[idx - 1] if idx - 1 < len(keywords) else "execution"
        enriched.append(f"• {bullet} (Aligned keyword: {kw})")
    return enriched


def _ats_refactor(original_text: str, job_desc: str) -> str:
    keywords = _extract_top_keywords(job_desc)
    resume_lines = [ln for ln in original_text.split("\n") if ln.strip()]

    summary_lines = [
        "ATS-Optimized Resume",
        "",
        "Professional Summary",
        "Results-driven professional with proven delivery and role-aligned impact.",
        f"Targeted ATS keywords: {', '.join(keywords[:10]) or 'N/A'}.",
        "",
        "Experience Highlights",
    ]

    summary_lines.extend(_build_bullets(resume_lines, keywords))
    summary_lines.extend(["", "Original Resume Content (preserved)", original_text.strip()])
    return "\n".join(summary_lines)


def _write_docx(content: str, template: Optional[Document], target: Path) -> None:
    doc = Document()
    sample = None
    if template and template.paragraphs and template.paragraphs[0].runs:
        sample = template.paragraphs[0].runs[0]

    for line in content.split("\n"):
        para = doc.add_paragraph(line)
        if sample and para.runs:
            run = para.runs[0]
            run.font.name = sample.font.name
            run.bold = sample.bold
            run.italic = sample.italic
            if sample.font.size:
                run.font.size = sample.font.size

    doc.save(target)


def _write_pdf(content: str, target: Path) -> None:
    c = canvas.Canvas(str(target), pagesize=letter)
    text_obj = c.beginText(40, 750)
    text_obj.setFont("Helvetica", 10)
    for line in content.split("\n"):
        text_obj.textLine(line[:120])
        if text_obj.getY() < 40:
            c.drawText(text_obj)
            c.showPage()
            text_obj = c.beginText(40, 750)
            text_obj.setFont("Helvetica", 10)
    c.drawText(text_obj)
    c.save()


def _write_rtf(content: str, target: Path) -> None:
    escaped = content.replace("\\", r"\\").replace("{", r"\{").replace("}", r"\}")
    target.write_text("{\\rtf1\\ansi\n" + escaped.replace("\n", "\\par\n") + "\n}", encoding="utf-8")


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/generate", response_model=GenerateResponse)
async def generate_resume(
    resume: UploadFile = File(...),
    job_url: str = Form(...),
    output_format: str = Form("all"),
) -> GenerateResponse:
    if output_format not in ALLOWED_FORMATS:
        raise HTTPException(status_code=400, detail=f"output_format must be one of: {sorted(ALLOWED_FORMATS)}")

    raw = await resume.read()
    text, template_doc = _read_resume(resume, raw)
    if not text.strip():
        raise HTTPException(status_code=400, detail="Resume appears to be empty.")

    job_title, company, job_desc = _extract_job(job_url)
    optimized = _ats_refactor(text, job_desc)

    date_tag = datetime.utcnow().strftime("%Y%m%d")
    id_tag = str(uuid4())[:8]
    stem = f"{_slugify(job_title)}_{date_tag}_{id_tag}"

    files: dict[str, str] = {}
    formats = {"docx", "pdf", "rtf"} if output_format == "all" else {output_format}

    if "docx" in formats:
        docx_path = OUTPUT_DIR / f"{stem}.docx"
        _write_docx(optimized, template_doc, docx_path)
        files["docx"] = f"/download/{docx_path.name}"

    if "pdf" in formats:
        pdf_path = OUTPUT_DIR / f"{stem}.pdf"
        _write_pdf(optimized, pdf_path)
        files["pdf"] = f"/download/{pdf_path.name}"

    if "rtf" in formats:
        rtf_path = OUTPUT_DIR / f"{stem}.rtf"
        _write_rtf(optimized, rtf_path)
        files["rtf"] = f"/download/{rtf_path.name}"

    preview = optimized[:2200]
    thumbnail = "\n".join(preview.splitlines()[:6])
    return GenerateResponse(job_title=job_title, company=company, preview=preview, thumbnail=thumbnail, files=files)


@app.get("/download/{filename}")
def download_file(filename: str):
    safe_name = Path(filename).name
    file_path = OUTPUT_DIR / safe_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path=file_path, filename=safe_name)
